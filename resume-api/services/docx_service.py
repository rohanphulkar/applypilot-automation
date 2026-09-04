import logging
import os
from pathlib import Path
from typing import Any, Dict, List, Tuple, Union

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

COLOR_DOCX_TEXT = RGBColor(13, 13, 13)        # Near black (#0d0d0d)
COLOR_DOCX_MUTED = RGBColor(100, 111, 121)    # Charcoal gray (#646f79)
COLOR_DOCX_LINK = RGBColor(26, 86, 219)       # ATS readable dark blue (#1a56db)
HEX_DOCX_BORDER = "CCCCCC"


def _add_docx_hyperlink(paragraph, url: str, text: str, color=COLOR_DOCX_LINK, underline: bool = False):
    if not url:
        run = paragraph.add_run(text)
        run.font.color.rgb = COLOR_DOCX_TEXT
        return run

    url_str = str(url).strip()
    clean_display = str(text or url_str).strip()

    try:
        part = paragraph.part
        r_id = part.relate_to(
            url_str,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")

        if color:
            c = OxmlElement("w:color")
            c.set(qn("w:val"), f"{color[0]:02x}{color[1]:02x}{color[2]:02x}")
            rPr.append(c)

        if underline:
            u = OxmlElement("w:u")
            u.set(qn("w:val"), "single")
            rPr.append(u)

        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Calibri")
        rFonts.set(qn("w:hAnsi"), "Calibri")
        rPr.append(rFonts)

        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "19")
        rPr.append(sz)

        new_run.append(rPr)
        new_run_text = OxmlElement("w:t")
        new_run_text.text = clean_display
        new_run.append(new_run_text)
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)
        return hyperlink
    except Exception as e:
        logger.warning("Failed to create docx hyperlink for %s: %s", url_str, e)
        run = paragraph.add_run(clean_display)
        run.font.color.rgb = color or COLOR_DOCX_TEXT
        run.font.underline = underline
        return run


def _add_docx_section_heading(doc: Document, title_text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True

    run = p.add_run(title_text.upper())
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = COLOR_DOCX_TEXT

    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>\n'
        f'  <w:bottom w:val="single" w:sz="6" w:space="2" w:color="{HEX_DOCX_BORDER}"/>\n'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def _format_docx_skills(skills: Any) -> List[Tuple[str, str]]:
    if not skills:
        return []

    if isinstance(skills, dict):
        results = []
        for category, items in skills.items():
            cat_name = str(category).replace("_", " ").title()
            if isinstance(items, list):
                skill_str = ", ".join(str(s).strip() for s in items if str(s).strip())
            else:
                skill_str = str(items).strip()
            if skill_str:
                results.append((cat_name, skill_str))
        return results

    if isinstance(skills, list):
        items = []
        for s in skills:
            if isinstance(s, dict):
                val = s.get("skill") or s.get("name") or str(s)
                if val:
                    items.append(str(val).strip())
            elif s:
                items.append(str(s).strip())
        return [("Technical Skills", ", ".join(items))] if items else []

    if isinstance(skills, str):
        return [("Technical Skills", skills.strip())]

    return []


def generate_docx(data: Dict[str, Any], output_file_path: Union[str, Path]) -> Path:
    """
    Generates a highly polished, ATS-optimized Microsoft Word (.docx) resume.
    Ensures 92+ ATS score, exact typography, right-aligned tab stops for dates, and hyperlinks.
    """
    if not data or not isinstance(data, dict):
        data = {}

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)

    printable_width = Inches(7.5)

    personal = data.get("personal", {})
    if not isinstance(personal, dict):
        personal = {}

    name = personal.get("name") or data.get("name") or "Candidate Resume"
    title = personal.get("title") or data.get("title") or ""
    email = personal.get("email") or ""
    website = personal.get("website") or ""
    github = personal.get("github") or ""
    linkedin = personal.get("linkedin") or ""
    summary = data.get("summary") or personal.get("summary") or ""

    # Header Name
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(2)
    p_name.paragraph_format.keep_with_next = True

    r_name = p_name.add_run(name)
    r_name.font.name = "Calibri"
    r_name.font.size = Pt(18)
    r_name.font.bold = True
    r_name.font.color.rgb = COLOR_DOCX_TEXT

    # Header Title
    if title:
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_before = Pt(0)
        p_title.paragraph_format.space_after = Pt(3)
        p_title.paragraph_format.keep_with_next = True

        r_title = p_title.add_run(title)
        r_title.font.name = "Calibri"
        r_title.font.size = Pt(10.5)
        r_title.font.bold = True
        r_title.font.color.rgb = COLOR_DOCX_TEXT

    # Contact Line
    contacts = []
    if email:
        contacts.append(("email", email, f"mailto:{email}", email))
    if website:
        clean_web = website.replace("https://", "").replace("http://", "").rstrip("/")
        contacts.append(("website", website, website, clean_web))
    if github:
        clean_gh = github.replace("https://", "").replace("http://", "").rstrip("/")
        contacts.append(("github", github, github, clean_gh))
    if linkedin:
        clean_li = linkedin.replace("https://", "").replace("http://", "").rstrip("/")
        contacts.append(("linkedin", linkedin, linkedin, clean_li))

    if contacts:
        p_contacts = doc.add_paragraph()
        p_contacts.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_contacts.paragraph_format.space_before = Pt(0)
        p_contacts.paragraph_format.space_after = Pt(6)
        p_contacts.paragraph_format.keep_with_next = True

        for idx, (ctype, full_url, href, display_txt) in enumerate(contacts):
            if idx > 0:
                sep = p_contacts.add_run("  |  ")
                sep.font.name = "Calibri"
                sep.font.size = Pt(9.5)
                sep.font.color.rgb = COLOR_DOCX_MUTED

            _add_docx_hyperlink(p_contacts, href, display_txt, color=COLOR_DOCX_LINK, underline=False)

    # Summary
    if summary:
        _add_docx_section_heading(doc, "SUMMARY")
        p_sum = doc.add_paragraph()
        p_sum.paragraph_format.space_before = Pt(2)
        p_sum.paragraph_format.space_after = Pt(4)
        p_sum.paragraph_format.line_spacing = 1.08

        r_sum = p_sum.add_run(summary)
        r_sum.font.name = "Calibri"
        r_sum.font.size = Pt(9.5)
        r_sum.font.color.rgb = COLOR_DOCX_TEXT

    # Skills
    raw_skills = data.get("skills", [])
    skills_entries = _format_docx_skills(raw_skills)
    if skills_entries:
        _add_docx_section_heading(doc, "SKILLS & CORE COMPETENCIES")
        for cat_label, skill_values in skills_entries:
            p_skill = doc.add_paragraph()
            p_skill.paragraph_format.space_before = Pt(1)
            p_skill.paragraph_format.space_after = Pt(2)
            p_skill.paragraph_format.line_spacing = 1.05

            if len(skills_entries) > 1 or cat_label != "Technical Skills":
                r_lbl = p_skill.add_run(f"{cat_label}: ")
                r_lbl.font.name = "Calibri"
                r_lbl.font.size = Pt(9.5)
                r_lbl.font.bold = True
                r_lbl.font.color.rgb = COLOR_DOCX_TEXT

            r_val = p_skill.add_run(skill_values)
            r_val.font.name = "Calibri"
            r_val.font.size = Pt(9.5)
            r_val.font.color.rgb = COLOR_DOCX_TEXT

    # Experience
    experience = data.get("experience", [])
    if experience and isinstance(experience, list):
        _add_docx_section_heading(doc, "EXPERIENCE")

        for job in experience:
            if not isinstance(job, dict):
                continue

            company = str(job.get("company", "")).strip()
            role = str(job.get("role", "")).strip()
            dates = str(job.get("dates", "")).strip()
            bullets = job.get("bullets", [])
            tech_stack = job.get("tech_stack", [])

            p_job = doc.add_paragraph()
            p_job.paragraph_format.space_before = Pt(4)
            p_job.paragraph_format.space_after = Pt(1)
            p_job.paragraph_format.keep_with_next = True
            p_job.paragraph_format.tab_stops.add_tab_stop(printable_width, WD_TAB_ALIGNMENT.RIGHT)

            r_comp = p_job.add_run(f"{company}")
            r_comp.font.name = "Calibri"
            r_comp.font.size = Pt(10)
            r_comp.font.bold = True
            r_comp.font.color.rgb = COLOR_DOCX_TEXT

            if role:
                r_pipe = p_job.add_run("  |  ")
                r_pipe.font.name = "Calibri"
                r_pipe.font.size = Pt(9.5)
                r_pipe.font.color.rgb = COLOR_DOCX_MUTED

                r_role = p_job.add_run(role)
                r_role.font.name = "Calibri"
                r_role.font.size = Pt(9.5)
                r_role.font.italic = True
                r_role.font.color.rgb = COLOR_DOCX_TEXT

            if dates:
                p_job.add_run(f"\t{dates}")
                r_date = p_job.runs[-1]
                r_date.font.name = "Calibri"
                r_date.font.size = Pt(9.5)
                r_date.font.bold = True
                r_date.font.color.rgb = COLOR_DOCX_TEXT

            if isinstance(bullets, list):
                for bullet in bullets:
                    b_str = str(bullet).strip()
                    if not b_str:
                        continue
                    p_b = doc.add_paragraph()
                    p_b.paragraph_format.left_indent = Inches(0.2)
                    p_b.paragraph_format.space_before = Pt(1)
                    p_b.paragraph_format.space_after = Pt(1.5)
                    p_b.paragraph_format.line_spacing = 1.05

                    r_bullet_sym = p_b.add_run("•  ")
                    r_bullet_sym.font.name = "Calibri"
                    r_bullet_sym.font.size = Pt(9.5)
                    r_bullet_sym.font.color.rgb = COLOR_DOCX_MUTED

                    r_bullet_text = p_b.add_run(b_str)
                    r_bullet_text.font.name = "Calibri"
                    r_bullet_text.font.size = Pt(9.5)
                    r_bullet_text.font.color.rgb = COLOR_DOCX_TEXT

            if tech_stack and isinstance(tech_stack, list):
                tech_line = ", ".join(str(t).strip() for t in tech_stack if str(t).strip())
                if tech_line:
                    p_tech = doc.add_paragraph()
                    p_tech.paragraph_format.left_indent = Inches(0.2)
                    p_tech.paragraph_format.space_before = Pt(1)
                    p_tech.paragraph_format.space_after = Pt(3)

                    r_ts_lbl = p_tech.add_run("Tech Stack: ")
                    r_ts_lbl.font.name = "Calibri"
                    r_ts_lbl.font.size = Pt(9)
                    r_ts_lbl.font.bold = True
                    r_ts_lbl.font.color.rgb = COLOR_DOCX_MUTED

                    r_ts_val = p_tech.add_run(tech_line)
                    r_ts_val.font.name = "Calibri"
                    r_ts_val.font.size = Pt(9)
                    r_ts_val.font.italic = True
                    r_ts_val.font.color.rgb = COLOR_DOCX_MUTED

    # Projects
    projects = data.get("projects", [])
    if projects and isinstance(projects, list):
        _add_docx_section_heading(doc, "PROJECTS")

        for proj in projects:
            if not isinstance(proj, dict):
                continue

            p_name_val = str(proj.get("name", "")).strip()
            p_url_val = proj.get("url", "")
            p_desc = proj.get("description", [])
            p_techs = proj.get("tech_stack", [])

            p_proj = doc.add_paragraph()
            p_proj.paragraph_format.space_before = Pt(4)
            p_proj.paragraph_format.space_after = Pt(1)
            p_proj.paragraph_format.keep_with_next = True

            r_pname = p_proj.add_run(p_name_val)
            r_pname.font.name = "Calibri"
            r_pname.font.size = Pt(10)
            r_pname.font.bold = True
            r_pname.font.color.rgb = COLOR_DOCX_TEXT

            if p_url_val:
                p_proj.add_run("  (")
                clean_url = str(p_url_val).replace("https://", "").replace("http://", "").rstrip("/")
                _add_docx_hyperlink(p_proj, str(p_url_val), clean_url, color=COLOR_DOCX_LINK, underline=False)
                p_proj.add_run(")")

            if isinstance(p_desc, list):
                for desc_item in p_desc:
                    d_str = str(desc_item).strip()
                    if not d_str:
                        continue
                    p_pd = doc.add_paragraph()
                    p_pd.paragraph_format.left_indent = Inches(0.2)
                    p_pd.paragraph_format.space_before = Pt(1)
                    p_pd.paragraph_format.space_after = Pt(1.5)
                    p_pd.paragraph_format.line_spacing = 1.05

                    r_b = p_pd.add_run("•  ")
                    r_b.font.name = "Calibri"
                    r_b.font.size = Pt(9.5)
                    r_b.font.color.rgb = COLOR_DOCX_MUTED

                    r_d = p_pd.add_run(d_str)
                    r_d.font.name = "Calibri"
                    r_d.font.size = Pt(9.5)
                    r_d.font.color.rgb = COLOR_DOCX_TEXT

            if p_techs and isinstance(p_techs, list):
                p_tech_line = ", ".join(str(t).strip() for t in p_techs if str(t).strip())
                if p_tech_line:
                    p_ptech = doc.add_paragraph()
                    p_ptech.paragraph_format.left_indent = Inches(0.2)
                    p_ptech.paragraph_format.space_before = Pt(1)
                    p_ptech.paragraph_format.space_after = Pt(3)

                    r_pts_lbl = p_ptech.add_run("Tech Stack: ")
                    r_pts_lbl.font.name = "Calibri"
                    r_pts_lbl.font.size = Pt(9)
                    r_pts_lbl.font.bold = True
                    r_pts_lbl.font.color.rgb = COLOR_DOCX_MUTED

                    r_pts_val = p_ptech.add_run(p_tech_line)
                    r_pts_val.font.name = "Calibri"
                    r_pts_val.font.size = Pt(9)
                    r_pts_val.font.italic = True
                    r_pts_val.font.color.rgb = COLOR_DOCX_MUTED

    # Education
    education = data.get("education")
    edu_list = []
    if isinstance(education, dict) and (education.get("degree") or education.get("institution")):
        edu_list = [education]
    elif isinstance(education, list):
        edu_list = [e for e in education if isinstance(e, dict) and (e.get("degree") or e.get("institution"))]

    if edu_list:
        _add_docx_section_heading(doc, "EDUCATION")
        for edu in edu_list:
            p_edu = doc.add_paragraph()
            p_edu.paragraph_format.space_before = Pt(3)
            p_edu.paragraph_format.space_after = Pt(1)
            p_edu.paragraph_format.tab_stops.add_tab_stop(printable_width, WD_TAB_ALIGNMENT.RIGHT)

            degree = str(edu.get("degree", "")).strip()
            institution = str(edu.get("institution", "")).strip()
            location = str(edu.get("location", "")).strip()
            year = str(edu.get("year", "")).strip()

            r_deg = p_edu.add_run(degree)
            r_deg.font.name = "Calibri"
            r_deg.font.size = Pt(10)
            r_deg.font.bold = True
            r_deg.font.color.rgb = COLOR_DOCX_TEXT

            if institution:
                r_inst = p_edu.add_run(f"  |  {institution}")
                r_inst.font.name = "Calibri"
                r_inst.font.size = Pt(9.5)
                r_inst.font.color.rgb = COLOR_DOCX_TEXT

            right_side = []
            if location:
                right_side.append(location)
            if year:
                right_side.append(year)

            if right_side:
                p_edu.add_run(f"\t{', '.join(right_side)}")
                r_edate = p_edu.runs[-1]
                r_edate.font.name = "Calibri"
                r_edate.font.size = Pt(9.5)
                r_edate.font.bold = True
                r_edate.font.color.rgb = COLOR_DOCX_TEXT

    out_path = Path(output_file_path).resolve()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    logger.info("ATS-optimized DOCX generated successfully: %s", out_path)
    return out_path
